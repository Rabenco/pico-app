import os
import fitz  # PyMuPDF
import pytesseract
import tempfile
from PIL import Image, ImageEnhance, ImageFilter
from transformers import AutoTokenizer, AutoModelForTokenClassification, pipeline
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment
from tqdm import tqdm
import torch
import re
import string
from difflib import SequenceMatcher

# 📁 Configuration
DOSSIER_PDF = "fichier_pdf_test"
FICHIER_WORD = os.path.join(DOSSIER_PDF, "rapport_PICO_BioELECTRA_stable.docx")
FICHIER_EXCEL = os.path.join(DOSSIER_PDF, "tableau_PICO.xlsx")
MODEL_NAME = "kamalkraj/BioELECTRA-PICO"
MAX_TOKENS = 512

# 🔍 Modèle
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
tokenizer = AutoTokenizer.from_pretrained(MODEL_NAME)
model = AutoModelForTokenClassification.from_pretrained(MODEL_NAME).to(device)
ner_pipeline = pipeline("ner", model=model, tokenizer=tokenizer, aggregation_strategy="simple", device=0 if torch.cuda.is_available() else -1)

# 📷 OCR fallback
pytesseract.pytesseract.tesseract_cmd = r"C:\\Tesseract\\tesseract.exe"

def extract_text_from_pdf(pdf_path):
    try:
        doc = fitz.open(pdf_path)
        text = "\n".join(page.get_text("text") for page in doc).strip()
        if len(text) > 500:
            return text
        ocr_text = []
        for page in doc:
            pix = page.get_pixmap(dpi=300)
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                img_path = tmp.name
                pix.save(img_path)
                img = Image.open(img_path).convert("L")
                img = img.filter(ImageFilter.SHARPEN)
                img = ImageEnhance.Contrast(img).enhance(2.0)
                text = pytesseract.image_to_string(img, lang='eng+fra', config='--psm 6')
                ocr_text.append(text)
                os.remove(img_path)
        return "\n".join(ocr_text).strip()
    except Exception as e:
        print(f"❌ Erreur d'extraction OCR : {e}")
        return ""

def extract_title_and_year(text):
    lines = text.splitlines()
    title = next((l.strip() for l in lines if 10 < len(l) < 200 and l[0].isupper()), "Non détecté")
    year_match = re.search(r"(19|20)\d{2}", text)
    year = year_match.group(0) if year_match else "Non détecté"
    return title, year

def extract_abstract(text):
    match = re.search(r"(abstract|ABSTRACT)[\s\S]{200,2000}?(introduction|INTRODUCTION|methods|METHODS)", text)
    return match.group(0) if match else text[:2000]

def chunk_text(text, max_tokens=MAX_TOKENS):
    tokens = tokenizer.tokenize(text)
    chunks = []

    # Boucle sur les tokens avec fenêtre glissante
    for i in range(0, len(tokens), max_tokens - 50):
        chunk_tokens = tokens[i:i + max_tokens - 2]  # On garde une marge de sécurité
        input_ids = tokenizer.convert_tokens_to_ids(chunk_tokens)
        chunk_text = tokenizer.decode(input_ids, skip_special_tokens=True)
        chunks.append(chunk_text)

    return chunks



def smart_clean(text):
    text = re.sub(r"\s+", " ", text).strip()
    text = re.sub(r"\b(?:group|study|trial|site|sites|patients|enrolled|assigned|control|device)\b", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\b\d+\b", "", text)
    text = re.sub(r"\(.*?\)", "", text)
    tokens = text.split()
    cleaned = []
    seen = set()
    for i in range(len(tokens)):
        current = tokens[i]
        if current.lower() in seen:
            continue
        duplicate = False
        for j in range(max(0, len(cleaned) - 6), len(cleaned)):
            ratio = SequenceMatcher(None, cleaned[j].lower(), current.lower()).ratio()
            if ratio > 0.95:
                duplicate = True
                break
        if not duplicate:
            cleaned.append(current)
            seen.add(current.lower())
    result = " ".join(cleaned)
    result = re.sub(r"\s+", " ", result)
    result = result.strip(string.punctuation + " ")
    return result[0].upper() + result[1:] if result else "Non détecté"

def extract_pico(text):
    grouped = {"Participants": [], "Intervention": [], "Comparator": [], "Outcome": []}
    for chunk in chunk_text(text):
        results = ner_pipeline(chunk)
        for entity in results:
            label = entity["entity_group"]
            if label in grouped:
                grouped[label].append(entity["word"])
    if not grouped["Comparator"]:
        i_text = " ".join(grouped["Intervention"])
        match = re.search(r"(vs\.?|versus|or)\s+(.+?)(,|;|$)", i_text, re.IGNORECASE)
        if match:
            grouped["Comparator"].append(match.group(2).strip())
    return {
        "P": smart_clean(" ".join(grouped["Participants"])),
        "I": smart_clean(" ".join(grouped["Intervention"])),
        "C": smart_clean(" ".join(grouped["Comparator"])),
        "O": smart_clean(" ".join(grouped["Outcome"]))
    }

def save_word(results, path):
    doc = Document()
    doc.add_heading("Extraction PICO avec BioELECTRA-PICO (stable)", 0)
    for res in results:
        doc.add_heading(res["Fichier"], level=1)
        for k in ["Titre", "Année", "P", "I", "C", "O"]:
            p = doc.add_paragraph()
            p.add_run(f"{k}: ").bold = True
            p.add_run(res[k])
    doc.save(path)
    print(f"📄 Rapport Word enregistré : {path}")

def save_excel(results, path):
    wb = Workbook()
    ws = wb.active
    ws.title = "PICO Résumé"
    headers = ["Fichier", "Titre", "Année", "P", "I", "C", "O"]
    ws.append(headers)
    for col in ws.iter_cols(min_row=1, max_row=1):
        for cell in col:
            cell.font = Font(bold=True)
    for res in results:
        ws.append([res[h] for h in headers])
    for col in ws.columns:
        col_letter = col[0].column_letter
        max_length = max(len(str(cell.value)) if cell.value else 0 for cell in col)
        ws.column_dimensions[col_letter].width = min(max_length + 2, 60)
        for cell in col:
            cell.alignment = Alignment(wrap_text=True)
    wb.save(path)
    print(f"📊 Tableau Excel enregistré : {path}")

def main():
    if not os.path.exists(DOSSIER_PDF):
        print("❌ Dossier introuvable.")
        return
    fichiers = [f for f in os.listdir(DOSSIER_PDF) if f.lower().endswith(".pdf")]
    if not fichiers:
        print("❌ Aucun PDF trouvé.")
        return
    résultats = []
    for fichier in tqdm(fichiers, desc="📊 Analyse des PDF"):
        chemin = os.path.join(DOSSIER_PDF, fichier)
        texte = extract_text_from_pdf(chemin)
        if not texte or len(texte) < 300:
            print(f"⚠️ Aucun texte exploitable : {fichier}")
            résultats.append({"Fichier": fichier, "Titre": "Non détecté", "Année": "Non détecté", "P": "", "I": "", "C": "", "O": ""})
            continue
        titre, annee = extract_title_and_year(texte)
        texte_cible = extract_abstract(texte)
        pico = extract_pico(texte_cible)
        résultats.append({"Fichier": fichier, "Titre": titre, "Année": annee, **pico})
    save_word(résultats, FICHIER_WORD)
    save_excel(résultats, FICHIER_EXCEL)

if __name__ == "__main__":
    main()
